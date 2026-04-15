#!/usr/bin/env python3
"""
上海理工大学博士/硕士研究生学位论文 Word 模板生成器
根据《上海理工大学博士、硕士研究生学位论文写作规范（2018年修订版）》生成
"""

from docx import Document
from docx.shared import Pt, Cm, Emu, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ============================================================
# 全局样式设置
# ============================================================

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)  # 小四号 = 12pt
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

pf = style.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
pf.line_spacing = Pt(20)  # 固定值20磅
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐

# ============================================================
# 页面设置：A4，页边距
# ============================================================
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(3.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(3.0)
section.header_distance = Cm(2.0)
section.footer_distance = Cm(2.0)

# 启用奇偶页不同页眉
section.different_first_page_header_footer = False
section.odd_and_even_pages_header_footer = True

# ============================================================
# 辅助函数
# ============================================================

def set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=False, color=None):
    """设置run的中英文字体和大小"""
    run.font.name = en_font
    run.font.size = size
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}></w:rPr>')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}></w:rFonts>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)

def add_chapter_title(text):
    """章标题：黑体小二号居中，段前24磅段后18磅，单倍行距"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(24)
    pf.space_after = Pt(18)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None
    run = p.add_run(text)
    set_run_font(run, cn_font='黑体', en_font='Times New Roman', size=Pt(18), bold=True)  # 小二号=18pt
    return p

def add_section1_title(text):
    """一级节标题：加粗宋体四号居左，段前24磅段后6磅"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(24)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(14), bold=True)  # 四号=14pt
    return p

def add_section2_title(text):
    """二级节标题：加粗宋体小四号居左，段前12磅段后6磅"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12), bold=True)  # 小四号=12pt
    return p

def add_body_text(text, first_line_indent=True):
    """正文：宋体小四号，首行缩进2字符，固定值20磅"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(20)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if first_line_indent:
        pf.first_line_indent = Cm(0.74)  # 约2个汉字符宽度
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12))
    return p

def add_figure_caption(text):
    """图名：宋体小四号居中，图下方，段前6磅段后12磅"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(6)
    pf.space_after = Pt(12)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12))
    return p

def add_table_caption(text):
    """表名：宋体小四号居中，表上方，段前12磅段后6磅"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = None
    run = p.add_run(text)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12))
    return p

def add_blank_lines(n=1):
    for _ in range(n):
        doc.add_paragraph()

# ============================================================
# 页眉页脚
# ============================================================

def setup_header_footer():
    """设置页眉页脚"""
    # 奇数页页眉（章标题）
    odd_header = section.header
    odd_header.is_linked_to_previous = False
    hp = odd_header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = hp.add_run('第一章  引言')  # 默认内容，需用户修改
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(10.5))  # 五号=10.5pt

    # 偶数页页眉（学校+论文类型）
    even_header = section.even_page_header
    even_header.is_linked_to_previous = False
    hp2 = even_header.paragraphs[0]
    hp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run2 = hp2.add_run('上海理工大学博士学位论文')
    set_run_font(run2, cn_font='宋体', en_font='Times New Roman', size=Pt(10.5))

    # 页脚（页码）
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    # 插入页码字段
    run_pg = fp.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run_pg._element.append(fldChar1)
    run_pg2 = fp.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run_pg2._element.append(instrText)
    run_pg3 = fp.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run_pg3._element.append(fldChar2)
    set_run_font(run_pg, size=Pt(10.5))
    set_run_font(run_pg2, size=Pt(10.5))
    set_run_font(run_pg3, size=Pt(10.5))

setup_header_footer()

# ============================================================
# 正文内容（模板结构）
# ============================================================

# ====== 中文摘要 ======
add_chapter_title('摘  要')

add_body_text(
    '在此撰写中文摘要。摘要应包括研究的目的、意义；论文的主要内容'
    '（作者独立进行的研究工作的概括性叙述）；获得的成果和结论。'
    '摘要字数应不少于500字。'
)
add_body_text(
    '（示例：本研究针对……的问题，采用……方法，通过……实验/分析，'
    '得出了……结论。本文的主要贡献在于……。）'
)
add_body_text(
    '（请替换以上示例文字为您的实际摘要内容。摘要部分的页眉为"摘要"。）'
)

# 关键词
p_kw = doc.add_paragraph()
p_kw.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_kw.paragraph_format.line_spacing = Pt(20)
p_kw.paragraph_format.space_before = Pt(12)
run_kw_label = p_kw.add_run('关键词：')
set_run_font(run_kw_label, cn_font='宋体', en_font='Times New Roman', size=Pt(14), bold=True)
run_kw = p_kw.add_run('关键词1  关键词2  关键词3  关键词4')
set_run_font(run_kw, cn_font='宋体', en_font='Times New Roman', size=Pt(14), bold=True)

# ====== 分页：英文摘要 ======
doc.add_page_break()

add_chapter_title('ABSTRACT')

add_body_text(
    'Write your English abstract here. The abstract should include the purpose '
    'and significance of the research, a summary of the main content and '
    'independent research work, and the results and conclusions obtained. '
    'The abstract should contain no fewer than 500 words.'
)
add_body_text(
    '(Replace this placeholder text with your actual English abstract.)'
)

# 英文关键词
p_kw_en = doc.add_paragraph()
p_kw_en.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_kw_en.paragraph_format.line_spacing = Pt(20)
p_kw_en.paragraph_format.space_before = Pt(12)
run_kw_en_label = p_kw_en.add_run('Key Words: ')
set_run_font(run_kw_en_label, cn_font='Times New Roman', en_font='Times New Roman', size=Pt(14), bold=True)
run_kw_en = p_kw_en.add_run('keyword1, keyword2, keyword3, keyword4')
set_run_font(run_kw_en, cn_font='Times New Roman', en_font='Times New Roman', size=Pt(14), bold=True)

# ====== 分页：目录 ======
doc.add_page_break()

add_chapter_title('目  录')

add_body_text('（此处插入自动目录。在Word中：引用 → 目录 → 自动目录，或手动填写。）', first_line_indent=False)

# 目录示例结构
toc_items = [
    ('中文摘要', '', False),
    ('ABSTRACT', '', False),
    ('第一章  引言', '.................................................. 1', False),
    ('  1.1  研究背景与意义', '..................................... 2', False),
    ('    1.1.1  国内外研究现状', '........................... 3', False),
    ('  1.2  研究内容与方法', '..................................... 5', False),
    ('第二章  ××××', '.................................................. 7', False),
    ('  2.1  ××××', '.................................................. 8', False),
    ('    2.1.1  ××××', '.......................................... 9', False),
]

for item, page, _ in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(item)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12))
    if page:
        run_pg = p.add_run(page)
        set_run_font(run_pg, cn_font='宋体', en_font='Times New Roman', size=Pt(12))

# ====== 分页：正文第一章 ======
doc.add_page_break()

add_chapter_title('第一章  引  言')

add_section1_title('1.1  研究背景与意义')

add_body_text(
    '在此撰写研究背景与意义。应主要论述论文的选题意义、国内外研究现状、'
    '本论文要解决的问题、论文运用的主要理论与方法、基本思路及论文的结构等。'
)

add_section2_title('1.1.1  国内外研究现状')

add_body_text(
    '在此撰写国内外研究现状。引用文献时使用上标编号，例如：目前，LaCrO₃的制备'
    '工艺主要有固态反应法[1]、共沉淀法[2]、溶胶-凝胶法[3]等。'
)

add_body_text(
    '（请在此处继续撰写相关内容。正文用宋体小四号字，行距固定值20磅。）'
)

add_section1_title('1.2  研究内容与方法')

add_body_text(
    '在此撰写研究内容与方法。利用前人或本人的理论和方法，解决别人没有做过的或'
    '没有解决的问题，应有一定的新见解或新内容，且概念清晰、分析严谨、推导正确。'
)

# ====== 分页：第二章 ======
doc.add_page_break()

add_chapter_title('第二章  ××××××××')

add_section1_title('2.1  ××××××')

add_body_text('在此撰写第二章一级节内容。')

add_section2_title('2.1.1  ××××××')

add_body_text('在此撰写第二章二级节内容。')

# ====== 图表示例 ======
add_blank_lines(1)
add_figure_caption('图 2.1  ××××示意图')

add_blank_lines(1)

# 表示例（三线表）
add_table_caption('表 2.1  ××××数据表')

table = doc.add_table(rows=4, cols=4, style='Table Grid')
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 设置三线表样式：只有顶线、底线、栏目线
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    '  <w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
    '  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
    '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>'
)
tblPr.append(borders)

# 填充表头
headers = ['序号', '指标名称', '数值', '单位']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = ''
    run = cell.paragraphs[0].add_run(h)
    set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(10.5), bold=True)  # 五号
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 填充数据行
data = [
    ['1', '×××', '××', '××'],
    ['2', '×××', '××', '××'],
    ['3', '×××', '××', '××'],
]
for r_idx, row_data in enumerate(data):
    for c_idx, val in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(10.5))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# ====== 分页：第三章 ======
doc.add_page_break()

add_chapter_title('第三章  ××××××××')

add_section1_title('3.1  ××××××')

add_body_text('在此撰写第三章内容。')

add_section2_title('3.1.1  ××××××')

add_body_text('在此撰写相关内容。')

# ====== 分页：第四章 ======
doc.add_page_break()

add_chapter_title('第四章  ××××××××')

add_section1_title('4.1  ××××××')

add_body_text('在此撰写第四章内容。')

# ====== 分页：结论 ======
doc.add_page_break()

add_chapter_title('结  论')

add_body_text(
    '在此撰写论文结论。论文结论要明确、精炼、完整、准确，突出自己的创造性成果'
    '或新见解。一般包括：本文的主要研究工作、主要创新点、研究不足与展望。'
)

add_section1_title('主要研究工作')

add_body_text('（总结本文完成的主要研究工作。）')

add_section1_title('主要创新点')

add_body_text('（列出本文的创新之处。）')

add_section1_title('研究不足与展望')

add_body_text('（分析研究的局限性和未来展望。）')

# ====== 分页：参考文献 ======
doc.add_page_break()

add_chapter_title('参考文献')

# 参考文献格式示例
refs = [
    '[1] XXX, XXX, XXX. 真空预冷过程中真空泵的延迟开启现象研究[J]. 制冷学报, 2011, (2):45-49.',
    '[2] Brauner N. Vapor absorption into falling film[J]. ASME J, 1991, 34(3): 76-82.',
    '[3] XXX. 出版集团研究[M]. 北京:中国书籍出版社, 2000:179-193.',
    '[4] Baehr H D. Heat and mass transfer[M]. Berlin: Springer-Verlag, 1994: 221.',
    '[5] XXX 等. 吸收式捕水器: 中国, 201210457660[P]. 2013-02-06.',
    '[6] GB/T7714—2015. 信息与文献 参考文献著录规则[S]. 北京:中国标准出版社, 2015.',
    '[7] XXX. 间断动力系统的随机扰动及其在守恒律方程中的应用[D]. 北京:北京大学数学学院, 1998.',
    '[8] Online Computer Library Center, Inc. History of OCLC[EB/OL]. [2000-01-08]. http://www.oclc.org/about/history/default.htm.',
    '[9] （请替换为您的实际参考文献）',
]

for ref in refs:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(16)  # 参考文献固定值16磅
    pf.space_before = Pt(3)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # 判断是否有中文
    has_cn = any('\u4e00' <= c <= '\u9fff' for c in ref)
    run = p.add_run(ref)
    if has_cn:
        set_run_font(run, cn_font='宋体', en_font='Times New Roman', size=Pt(12))
    else:
        set_run_font(run, cn_font='Times New Roman', en_font='Times New Roman', size=Pt(12))

# ====== 分页：附录 ======
doc.add_page_break()

add_chapter_title('附录A  ××××××')

add_body_text('在此撰写附录内容。附录格式与正文相同。')
add_body_text('附录中的图、表、数学表达式另行编序号，冠以附录序号，如"图A.1"、"表A.1"。')

# ====== 分页：在读期间成果 ======
doc.add_page_break()

add_chapter_title('在读期间公开发表的论文和承担科研项目及取得成果')

p_pub = doc.add_paragraph()
p_pub.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_pub.paragraph_format.line_spacing = Pt(20)
p_pub.paragraph_format.space_before = Pt(24)
p_pub.paragraph_format.space_after = Pt(6)
run_pub = p_pub.add_run('一、论文')
set_run_font(run_pub, cn_font='宋体', en_font='Times New Roman', size=Pt(14), bold=True)

add_body_text(
    '1. XXX 等. 文章标题[J]. 刊名, 出版年, 卷(期):页码.', first_line_indent=False
)
add_body_text(
    '2. XXX, XXX. 文章标题[J]. 刊名, 出版年, 卷(期):页码.', first_line_indent=False
)

p_pub2 = doc.add_paragraph()
p_pub2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_pub2.paragraph_format.line_spacing = Pt(20)
p_pub2.paragraph_format.space_before = Pt(24)
p_pub2.paragraph_format.space_after = Pt(6)
run_pub2 = p_pub2.add_run('二、专利')
set_run_font(run_pub2, cn_font='宋体', en_font='Times New Roman', size=Pt(14), bold=True)

add_body_text(
    '1. 发明专利: 专利名称, 专利号: xxxxx, 第x发明人.', first_line_indent=False
)

p_pub3 = doc.add_paragraph()
p_pub3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
p_pub3.paragraph_format.line_spacing = Pt(20)
p_pub3.paragraph_format.space_before = Pt(24)
p_pub3.paragraph_format.space_after = Pt(6)
run_pub3 = p_pub3.add_run('三、科研项目')
set_run_font(run_pub3, cn_font='宋体', en_font='Times New Roman', size=Pt(14), bold=True)

add_body_text(
    '1. 项目来源, 项目名称(项目编号), 起止时间, 参与级别.', first_line_indent=False
)

# ====== 分页：致谢 ======
doc.add_page_break()

add_chapter_title('致  谢')

add_body_text(
    '在此撰写致谢内容。主要向导师等对论文工作有直接贡献和帮助的人士、单位表示感谢。'
    '致谢应谦虚诚恳，实事求是，切忌浮夸及庸俗。致谢篇幅一般不超过一页。'
)
add_body_text(
    '（请替换为您的实际致谢内容。）'
)

# ============================================================
# 封面（使用分节符放在最前面 - 这里在文档末尾说明，实际使用时移到最前）
# ============================================================

# 我们在文档开头插入封面
# 先保存当前分节数，然后在开头插入新节

# 获取第一个section并添加封面内容
# 由于python-docx的限制，我们在文档最前面已经建立了内容
# 这里用bookmark或直接说明：封面由学院统一发放，内封面可从研究生院网站下载

# ============================================================
# 保存
# ============================================================

output_path = '/root/.openclaw/workspace/上海理工大学学位论文模板.docx'
doc.save(output_path)
print(f'模板已生成: {output_path}')
